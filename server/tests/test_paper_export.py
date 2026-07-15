from __future__ import annotations

import base64
import asyncio
import hashlib
import io
import json
import zipfile
from copy import deepcopy
from itertools import product
from pathlib import Path

import pytest
from fastapi.responses import StreamingResponse
from PIL import Image
from pypdf import PdfReader
from sqlalchemy import select
from starlette.requests import Request

from server.tests.test_teacher import teacher_context
from server.tests.test_teacher_papers import create_manual_paper, csrf, seed_question


FIXTURE = json.loads((Path(__file__).parent / "fixtures" / "export_expected.json").read_text())
PROJECT_ROOT = Path(__file__).resolve().parents[2]
REAL_MANIFEST = PROJECT_ROOT / "content" / "question-banks" / "soil-mechanics" / "manifest.json"
PNG_DATA_URI = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAADAAAAAgCAIAAADbtmxLAAAAaUlEQVR4nO3XsQ3DQAzAQD7hebKK50iR"
    "KTKCCy+VKhNlhHeRggb+agEi1Gns7w8lEiMxEiMxEiMxErNdGTpfj78sex7f+11IYiRGYiRGYiRGYiRG"
    "YiRGYiRGYiRmrL9sYgXNrKCZFcTED+GoBn1Ws7ToAAAAAElFTkSuQmCC"
)


def png_bytes(color: tuple[int, int, int] = (12, 90, 160), size: tuple[int, int] = (48, 32)) -> bytes:
    output = io.BytesIO()
    Image.new("RGB", size, color=color).save(output, format="PNG")
    return output.getvalue()


def png_data_uri(color: tuple[int, int, int]) -> str:
    return "data:image/png;base64," + base64.b64encode(png_bytes(color)).decode("ascii")


def paper_with_question(question: dict, *, title: str = "导出安全测试", duration: int | None = 90) -> dict:
    return {
        "title": title,
        "subject": {"name": "土力学"},
        "durationMinutes": duration,
        "totalPoints": 10,
        "paperQuestions": [
            {
                "sectionTitle": "一、测试题",
                "sequence": 1,
                "points": 10,
                "question": question,
            }
        ],
    }


def docx_media_count(content: bytes) -> int:
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        return sum(name.startswith("word/media/") for name in archive.namelist())


def pdf_image_count(content: bytes) -> int:
    count = 0
    for page in PdfReader(io.BytesIO(content)).pages:
        resources = page["/Resources"].get_object()
        xobjects = resources.get("/XObject")
        if xobjects is None:
            continue
        for reference in xobjects.get_object().values():
            if reference.get_object().get("/Subtype") == "/Image":
                count += 1
    return count


def media_count(content: bytes, format: str) -> int:
    return docx_media_count(content) if format == "docx" else pdf_image_count(content)


def document_order_text(content: bytes, format: str) -> str:
    if format == "pdf":
        return extract_pdf_text(content)
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        return archive.read("word/document.xml").decode("utf-8")


@pytest.fixture()
def sample_paper() -> dict:
    return {
        "title": FIXTURE["title"],
        "subject": {"name": FIXTURE["subject"]},
        "durationMinutes": 90,
        "totalPoints": 30,
        "paperQuestions": [
            {
                "sectionTitle": "一、选择题",
                "sequence": 1,
                "points": 5,
                "question": {
                    "text": "达西定律适用于下列哪种渗流状态？",
                    "questionType": "单项选择题",
                    "options": [{"label": "A", "text": "层流"}, {"label": "B", "text": "紊流"}],
                    "correctAnswer": "A",
                    "explanation": "达西定律描述层流渗透。",
                    "attachments": [{"kind": "image", "url": PNG_DATA_URI, "alt": "渗流示意图"}],
                    "sourceMetadata": {"formulas": ["q = k A i"]},
                },
            },
            {
                "sectionTitle": "二、综合题",
                "sequence": 2,
                "points": 25,
                "question": {
                    "text": "说明渗透系数的物理意义，并根据下表计算渗流量。",
                    "questionType": "计算题",
                    "options": [],
                    "correctAnswer": {"value": "0.002 m3/s"},
                    "explanation": "先确定水力坡降，再代入 q = kAi。",
                    "rubric": [{"criterion": "列出公式", "points": 10}, {"criterion": "计算结果", "points": 15}],
                    "attachments": [{"kind": "table", "rows": [["参数", "数值"], ["k", "1e-4 m/s"], ["A", "20 m2"]]}],
                    "sourceMetadata": {"tables": [[["参数", "数值"], ["i", "0.01"]]], "formulas": ["q = k A i"]},
                },
            },
        ],
    }


@pytest.fixture()
def ordered_attachment_paper() -> dict:
    stem_position = {"blockIndex": 10, "paragraphIndex": 10}
    answer_position = {"blockIndex": 11, "paragraphIndex": 11}

    def attachment(kind: str, ordinal: int, position: dict, **values) -> dict:
        return {
            "kind": kind,
            "inlineOrdinal": ordinal,
            "placeholder": f"[[attachment:{ordinal}]]",
            "sourcePosition": position,
            **values,
        }

    question = {
        "text": "公开题干",
        "questionType": "计算题",
        "options": [],
        "correctAnswer": {"value": 12.5, "unit": "kPa"},
        "explanation": "公开解析",
        "attachments": [
            attachment("formula", 1, stem_position, ommlText="FORMULA_STEM", ommlSource="<m:oMath>FORMULA_STEM</m:oMath>"),
            attachment("table", 2, stem_position, rows=[["TABLE_STEM", "1"]]),
            attachment("image", 3, stem_position, data=png_data_uri((20, 110, 180)), alt="IMAGE_STEM"),
            attachment("formula", 1, answer_position, ommlText="FORMULA_SECRET", ommlSource="<m:oMath>FORMULA_SECRET</m:oMath>"),
            attachment("table", 2, answer_position, rows=[["TABLE_SECRET", "2"]]),
            attachment("image", 3, answer_position, data=png_data_uri((180, 40, 60)), alt="IMAGE_SECRET"),
        ],
        "sourceMetadata": {
            "answerPosition": answer_position,
            "sourceBlocks": [
                {
                    "kind": "paragraph",
                    "role": "stem",
                    "sourcePosition": stem_position,
                    "text": "公开题干",
                    "textWithPlaceholders": "公开题干[[attachment:1]][[attachment:2]][[attachment:3]]",
                    "inlineContent": [
                        {"kind": "text", "text": "公开题干"},
                        {"kind": "attachment", "attachmentOrdinal": 1, "placeholder": "[[attachment:1]]"},
                        {"kind": "attachment", "attachmentOrdinal": 2, "placeholder": "[[attachment:2]]"},
                        {"kind": "attachment", "attachmentOrdinal": 3, "placeholder": "[[attachment:3]]"},
                    ],
                },
                {
                    "kind": "paragraph",
                    "role": "answer",
                    "sourcePosition": answer_position,
                    "text": "答案侧内容",
                    "textWithPlaceholders": "答案侧内容[[attachment:1]][[attachment:2]][[attachment:3]]",
                    "inlineContent": [
                        {"kind": "text", "text": "答案侧内容"},
                        {"kind": "attachment", "attachmentOrdinal": 1, "placeholder": "[[attachment:1]]"},
                        {"kind": "attachment", "attachmentOrdinal": 2, "placeholder": "[[attachment:2]]"},
                        {"kind": "attachment", "attachmentOrdinal": 3, "placeholder": "[[attachment:3]]"},
                    ],
                },
            ],
        },
    }
    return paper_with_question(question)


def extract_docx_text(content: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(content))
    return "\n".join(paragraph.text for paragraph in document.paragraphs) + "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )


def extract_pdf_text(content: bytes) -> str:
    return "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(content)).pages)


@pytest.mark.parametrize("format", ["docx", "pdf"])
def test_question_exports_include_structured_question_content_without_answers(sample_paper, format: str) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    content = render_paper(sample_paper, "questions", format, ExportOptions())
    text = extract_docx_text(content) if format == "docx" else extract_pdf_text(content)
    assert FIXTURE["title"] in text
    assert FIXTURE["duration"] in text
    assert FIXTURE["total"] in text
    assert "一、选择题" in text
    assert "二、综合题" in text
    assert "A. 层流" in text
    assert "q = k A i" in text
    assert FIXTURE["answer_heading"] not in text
    assert FIXTURE["correct_answer"] not in text
    assert "0.002 m3/s" not in text


def test_docx_question_export_has_embedded_media_and_tables(sample_paper) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    content = render_paper(sample_paper, "questions", "docx", ExportOptions())
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml")
        styles_xml = archive.read("word/styles.xml")
        settings_xml = archive.read("word/settings.xml")
    assert any(name.startswith("word/media/") for name in names)
    assert b"<w:tbl" in document_xml
    assert b"<w:hyperlink" not in document_xml
    assert b'cx="1800000"' in document_xml
    for slot in (b"ascii", b"hAnsi", b"eastAsia", b"cs"):
        assert slot + b'="Arial Unicode MS"' in document_xml + styles_xml
    assert b'w:hint="eastAsia"' in document_xml
    assert b'w:eastAsia="zh-CN"' in document_xml + styles_xml
    assert b'<w:themeFontLang w:val="zh-CN" w:eastAsia="zh-CN"' in settings_xml


def test_pdf_question_export_embeds_a_chinese_font_for_portable_rendering(sample_paper) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    content = render_paper(sample_paper, "questions", "pdf", ExportOptions())
    assert b"/FontFile2" in content


def test_answer_sheet_and_answers_variants_are_intentionally_distinct(sample_paper) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    answer_sheet = extract_docx_text(
        render_paper(sample_paper, "answer-sheet", "docx", ExportOptions())
    )
    answers = extract_docx_text(render_paper(sample_paper, "answers", "docx", ExportOptions()))
    assert answer_sheet.count(FIXTURE["answer_sheet_heading"]) == 1
    assert "正确答案：A" not in answer_sheet
    assert FIXTURE["answer_heading"] in answers
    assert FIXTURE["correct_answer"] in answers
    assert "0.002 m3/s" in answers


def test_export_endpoint_requires_owner_streams_rfc5987_filename_and_audits(teacher_context) -> None:
    from server.application.models import OperationLog
    from server.tests.test_teacher_catalog import teacher2_client

    client, database, _ = teacher_context
    paper = create_manual_paper(client, [seed_question(database, "export-owned-question")])
    response = client.get(
        f"/api/teacher/papers/{paper['id']}/export?format=pdf&variant=answers",
        headers=csrf(),
    )
    assert response.status_code == 200, response.text
    assert response.headers["content-type"].startswith("application/pdf")
    assert "filename*=UTF-8''" in response.headers["content-disposition"]
    assert "%E5%9C%9F" in response.headers["content-disposition"]
    assert extract_pdf_text(response.content).count("参考答案") == 1

    other = teacher2_client(client, database)
    denied = other.get(
        f"/api/teacher/papers/{paper['id']}/export?format=docx&variant=questions",
        headers=csrf("teacher-2-csrf"),
    )
    assert denied.status_code == 404
    assert denied.json()["code"] == "PAPER_NOT_FOUND"

    with database.session() as session:
        logs = session.scalars(
            select(OperationLog).where(
                OperationLog.action == "paper.export", OperationLog.target_id == paper["id"]
            )
        ).all()
    assert len(logs) == 1
    assert logs[0].detail == {"format": "pdf", "variant": "answers"}


@pytest.mark.parametrize(("variant", "format"), list(product(
    ["questions", "answer-sheet", "answers"], ["docx", "pdf"]
)))
def test_all_six_variant_format_combinations_are_complete_and_immutable(
    sample_paper, variant: str, format: str
) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    before = deepcopy(sample_paper)
    content = render_paper(sample_paper, variant, format, ExportOptions())
    text = extract_docx_text(content) if format == "docx" else extract_pdf_text(content)

    assert content.startswith(b"PK") if format == "docx" else content.startswith(b"%PDF")
    assert sample_paper == before
    assert FIXTURE["title"] in text
    if variant == "answer-sheet":
        assert FIXTURE["answer_sheet_heading"] in text
        assert "达西定律" not in text
        assert FIXTURE["correct_answer"] not in text
    else:
        assert "达西定律" in text
        assert (FIXTURE["answer_heading"] in text) is (variant == "answers")
        assert (FIXTURE["correct_answer"] in text) is (variant == "answers")


@pytest.mark.parametrize(("variant", "format"), list(product(
    ["questions", "answer-sheet", "answers"], ["docx", "pdf"]
)))
def test_source_blocks_preserve_attachment_order_and_never_leak_answer_side(
    ordered_attachment_paper, variant: str, format: str
) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    content = render_paper(ordered_attachment_paper, variant, format, ExportOptions())
    text = extract_docx_text(content) if format == "docx" else extract_pdf_text(content)
    ordered_text = document_order_text(content, format)

    if variant == "answer-sheet":
        assert media_count(content, format) == 0
        assert all(marker not in text for marker in ("FORMULA_STEM", "TABLE_STEM", "IMAGE_STEM"))
        assert all(marker not in text for marker in ("答案侧内容", "FORMULA_SECRET", "TABLE_SECRET", "IMAGE_SECRET"))
        return

    assert media_count(content, format) == (2 if variant == "answers" else 1)
    assert all(marker in text for marker in ("FORMULA_STEM", "TABLE_STEM", "IMAGE_STEM"))
    assert ordered_text.index("FORMULA_STEM") < ordered_text.index("TABLE_STEM") < ordered_text.index("IMAGE_STEM")
    assert "答案侧内容" not in text
    for marker in ("FORMULA_SECRET", "TABLE_SECRET", "IMAGE_SECRET"):
        assert (marker in text) is (variant == "answers")
    if variant == "answers":
        assert ordered_text.index("FORMULA_SECRET") < ordered_text.index("TABLE_SECRET") < ordered_text.index("IMAGE_SECRET")


@pytest.mark.parametrize("format", ["docx", "pdf"])
def test_structured_answers_are_localized_without_losing_fields(format: str) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    values = [
        ("判断题", False),
        ("多项选择题", ["A", "C"]),
        ("填空题", ["孔隙比", "空隙比"]),
        ("计算题", {"value": 12.5, "unit": "cm3/s", "tolerance": 0.1}),
    ]
    paper = {
        "title": "结构化答案测试",
        "durationMinutes": 30,
        "totalPoints": 40,
        "paperQuestions": [
            {
                "sectionTitle": "一、答案",
                "sequence": index,
                "points": 10,
                "question": {
                    "text": f"第 {index} 题",
                    "questionType": question_type,
                    "options": [],
                    "correctAnswer": answer,
                    "attachments": [],
                },
            }
            for index, (question_type, answer) in enumerate(values, start=1)
        ],
    }

    content = render_paper(paper, "answers", format, ExportOptions())
    text = extract_docx_text(content) if format == "docx" else extract_pdf_text(content)
    assert "正确答案：错误" in text
    assert "正确答案：A、C" in text
    assert "正确答案：可接受答案：孔隙比、空隙比" in text
    assert "正确答案：12.5 cm3/s；容差：0.1" in text
    assert "False" not in text


@pytest.mark.parametrize("format", ["docx", "pdf"])
def test_normalized_answers_do_not_repeat_imported_answer_or_explanation_text(format: str) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    answer_position = {"blockIndex": 20, "paragraphIndex": 20}
    explanation_position = {"blockIndex": 21, "paragraphIndex": 21}
    question = {
        "text": "达西定律适用于何种流态？",
        "questionType": "单项选择题",
        "options": [{"label": "A", "text": "紊流"}, {"label": "B", "text": "层流"}],
        "correctAnswer": "B",
        "explanation": "达西定律描述层流渗透。",
        "attachments": [],
        "sourceMetadata": {
            "answerPosition": answer_position,
            "sourceBlocks": [
                {"kind": "paragraph", "role": "answer", "sourcePosition": answer_position, "text": "答案：B"},
                {"kind": "paragraph", "role": "postAnswer", "sourcePosition": explanation_position, "text": "达西定律描述层流渗透。"},
            ],
        },
    }

    content = render_paper(paper_with_question(question), "answers", format, ExportOptions())
    text = extract_docx_text(content) if format == "docx" else extract_pdf_text(content)

    assert text.count("答案：B") == 1
    assert text.count("达西定律描述层流渗透。") == 1


@pytest.mark.parametrize("format", ["docx", "pdf"])
def test_real_imported_answer_block_is_not_duplicated(format: str) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    manifest = json.loads(REAL_MANIFEST.read_text(encoding="utf-8"))
    question = deepcopy(manifest["questions"][0])
    assert question["correctAnswer"] == "B"
    assert any(block.get("role") == "answer" for block in question["sourceBlocks"])

    content = render_paper(
        paper_with_question(question, title="真实导入答案去重"),
        "answers",
        format,
        ExportOptions(asset_root=PROJECT_ROOT / "public"),
    )
    text = extract_docx_text(content) if format == "docx" else extract_pdf_text(content)

    assert text.count("答案：B") == 1


@pytest.mark.parametrize("format", ["docx", "pdf"])
def test_nullable_duration_has_intentional_label(format: str) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    paper = paper_with_question(
        {"text": "无时限题", "questionType": "简答题", "options": [], "attachments": []},
        duration=None,
    )
    content = render_paper(paper, "questions", format, ExportOptions())
    text = extract_docx_text(content) if format == "docx" else extract_pdf_text(content)
    assert "考试时间：未限定" in text
    assert "None 分钟" not in text
    assert "未限定 分钟" not in text


@pytest.mark.parametrize("format", ["docx", "pdf"])
@pytest.mark.parametrize("source", ["/etc/passwd", "../secret.png", "question-assets/soil-mechanics/secret.png"])
def test_export_rejects_arbitrary_filesystem_image_sources(source: str, format: str) -> None:
    from server.application.services.paper_export import ExportOptions, PaperExportError, render_paper

    question = {
        "text": "不可信附件",
        "questionType": "简答题",
        "options": [],
        "attachments": [{"kind": "image", "path": source, "alt": "不应读取"}],
    }
    with pytest.raises(PaperExportError, match="图片"):
        render_paper(paper_with_question(question), "questions", format, ExportOptions(asset_root=PROJECT_ROOT / "public"))


def test_asset_resolver_enforces_allowlist_containment_digest_type_and_size(tmp_path: Path, monkeypatch) -> None:
    from server.application.services import paper_export
    from server.application.services.paper_export import ExportOptions, PaperExportError, render_paper

    root = tmp_path / "public"
    assets = root / "question-assets" / "soil-mechanics"
    assets.mkdir(parents=True)
    image = png_bytes()
    digest = hashlib.sha256(image).hexdigest()
    valid = assets / f"{digest}.png"
    valid.write_bytes(image)
    attachment = {
        "kind": "image",
        "src": f"/foundation-smart-companion/question-assets/soil-mechanics/{valid.name}",
        "sha256": digest,
        "alt": "可信资产",
    }
    question = {"text": "可信附件", "questionType": "简答题", "options": [], "attachments": [attachment]}
    content = render_paper(paper_with_question(question), "questions", "docx", ExportOptions(asset_root=root))
    assert docx_media_count(content) == 1

    missing = deepcopy(attachment)
    missing["src"] = "/foundation-smart-companion/question-assets/soil-mechanics/" + "0" * 64 + ".png"
    missing["sha256"] = "0" * 64
    with pytest.raises(PaperExportError, match="不存在"):
        render_paper(paper_with_question({**question, "attachments": [missing]}), "questions", "docx", ExportOptions(asset_root=root))

    invalid_digest = deepcopy(attachment)
    invalid_digest["sha256"] = "0" * 64
    with pytest.raises(PaperExportError, match="SHA-256"):
        render_paper(paper_with_question({**question, "attachments": [invalid_digest]}), "questions", "docx", ExportOptions(asset_root=root))

    wrong_mime = {"kind": "image", "data": "data:image/jpeg;base64," + base64.b64encode(image).decode("ascii")}
    with pytest.raises(PaperExportError, match="类型"):
        render_paper(paper_with_question({**question, "attachments": [wrong_mime]}), "questions", "docx", ExportOptions(asset_root=root))

    monkeypatch.setattr(paper_export, "MAX_IMAGE_BYTES", 32)
    oversized = {"kind": "image", "data": "data:image/png;base64," + base64.b64encode(image).decode("ascii")}
    with pytest.raises(PaperExportError, match="大小"):
        render_paper(paper_with_question({**question, "attachments": [oversized]}), "questions", "docx", ExportOptions(asset_root=root))

    outside = tmp_path / "outside.png"
    outside.write_bytes(image)
    symlink = assets / f"{hashlib.sha256(image).hexdigest()}.jpg"
    symlink.symlink_to(outside)
    escaped = deepcopy(attachment)
    escaped.update({"src": f"/foundation-smart-companion/question-assets/soil-mechanics/{symlink.name}"})
    with pytest.raises(PaperExportError, match="目录"):
        render_paper(paper_with_question({**question, "attachments": [escaped]}), "questions", "docx", ExportOptions(asset_root=root))


@pytest.mark.parametrize("format", ["docx", "pdf"])
def test_real_imported_png_asset_is_resolved_and_embedded_without_mutation(format: str) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    manifest = json.loads(REAL_MANIFEST.read_text(encoding="utf-8"))
    original = deepcopy(manifest)
    selected = None
    for question in manifest["questions"]:
        answer_index = question["sourceMetadata"].get("answerPosition", {}).get("blockIndex", 10**9)
        for attachment in question["attachments"]:
            if (
                attachment.get("kind") == "image"
                and attachment.get("src", "").endswith(".png")
                and attachment["sourcePosition"]["blockIndex"] < answer_index
            ):
                selected = (question, attachment)
                break
        if selected:
            break
    assert selected is not None
    question, attachment = deepcopy(selected)
    question["attachments"] = [attachment]
    paper = paper_with_question(question, title="真实题库图片导出")

    content = render_paper(paper, "questions", format, ExportOptions(asset_root=PROJECT_ROOT / "public"))

    assert media_count(content, format) >= 1
    assert json.loads(REAL_MANIFEST.read_text(encoding="utf-8")) == original


@pytest.mark.parametrize("format", ["docx", "pdf"])
def test_multi_page_and_oversized_table_rows_split_with_repeated_headers(format: str) -> None:
    from server.application.services.paper_export import ExportOptions, render_paper

    rows = [["编号", "说明"]]
    rows.extend([[str(index), f"第 {index} 行土力学参数说明"] for index in range(1, 75)])
    rows.append(["超长单元格", "可分页内容 " * 900])
    question = {
        "text": "分析下列多页试验记录并说明规律。" + "长题干 " * 160,
        "questionType": "计算题",
        "options": [],
        "attachments": [{"kind": "table", "rows": rows}],
    }

    content = render_paper(paper_with_question(question), "questions", format, ExportOptions())

    if format == "pdf":
        reader = PdfReader(io.BytesIO(content))
        assert len(reader.pages) >= 3
        page_texts = [page.extract_text() or "" for page in reader.pages]
        assert "分析下列多页试验记录" in page_texts[0]
        assert all(len(text.strip()) > 100 for text in page_texts)
        assert all("编号" in text for text in page_texts)
        assert "超长单元格（续" in "\n".join(page_texts)
        assert "超长单元格" in extract_pdf_text(content)
    else:
        from docx import Document

        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            xml = archive.read("word/document.xml")
        assert b"<w:tblHeader" in xml
        document = Document(io.BytesIO(content))
        oversized_row = next(
            row
            for table in document.tables
            for row in table.rows
            if row.cells[0].text == "超长单元格"
        )
        assert oversized_row._tr.xpath("./w:trPr/w:cantSplit")
        continuation_rows = [
            row
            for table in document.tables
            for row in table.rows
            if row.cells[0].text.startswith("超长单元格（续")
        ]
        assert continuation_rows
        assert all(row._tr.xpath("./w:trPr/w:cantSplit") for row in continuation_rows)
        assert "超长单元格" in extract_docx_text(content)


def test_endpoint_returns_streaming_response_and_encodes_every_filename_character(teacher_context) -> None:
    from server.application.api.dependencies import AuthContext
    from server.application.api.teacher_papers import export_paper
    from server.application.models import SessionToken, User

    client, database, _ = teacher_context
    title = "土力学 / 100% '重点'; 单元"
    paper = create_manual_paper(
        client,
        [seed_question(database, "export-stream-question")],
        title=title,
    )
    with database.session() as session:
        user = session.get(User, "teacher-user-1")
        token = session.get(SessionToken, "teacher-session")
        session.expunge(user)
        session.expunge(token)
    request = Request({"type": "http", "method": "GET", "path": "/", "headers": [], "app": client.app})

    response = export_paper(
        paper["id"], request, format="docx", variant="questions", auth=AuthContext(user=user, session=token)
    )

    assert isinstance(response, StreamingResponse)
    disposition = response.headers["content-disposition"]
    assert "%2F" in disposition
    assert "%25" in disposition
    assert "%27" in disposition
    assert "%3B" in disposition
    assert "%20" in disposition
    assert "filename*=UTF-8''" in disposition
    asyncio.run(response.body_iterator.aclose())


def test_all_six_endpoints_leave_paper_question_and_publication_snapshots_immutable(teacher_context) -> None:
    from server.application.models import AssignmentQuestion, Paper, Question

    client, database, _ = teacher_context
    question_id = seed_question(database, "export-immutable-question")
    paper = create_manual_paper(client, [question_id], title="不可变导出测试")
    published = client.post(
        f"/api/teacher/papers/{paper['id']}/publish",
        json={"studentIds": ["student-1"], "classIds": []},
        headers=csrf(),
    )
    assert published.status_code == 200, published.text
    assignment_id = published.json()["data"]["assignmentId"]

    def state() -> dict:
        with database.session() as session:
            stored_paper = session.get(Paper, paper["id"])
            stored_question = session.get(Question, question_id)
            snapshot = session.scalar(
                select(AssignmentQuestion).where(AssignmentQuestion.assignment_id == assignment_id)
            )
            return {
                "paper": {
                    "title": stored_paper.title,
                    "status": stored_paper.status,
                    "version": stored_paper.version,
                    "duration": stored_paper.duration_minutes,
                    "total": stored_paper.total_points,
                    "blueprint": deepcopy(stored_paper.assembly_blueprint),
                },
                "question": {
                    "text": stored_question.text,
                    "answer": deepcopy(stored_question.correct_answer),
                    "attachments": deepcopy(stored_question.attachments),
                    "metadata": deepcopy(stored_question.source_metadata),
                },
                "snapshot": deepcopy(snapshot.question_snapshot),
            }

    before = state()
    for variant, format in product(["questions", "answer-sheet", "answers"], ["docx", "pdf"]):
        response = client.get(
            f"/api/teacher/papers/{paper['id']}/export?variant={variant}&format={format}",
            headers=csrf(),
        )
        assert response.status_code == 200, response.text
    assert state() == before


def test_endpoint_rejects_unsafe_required_image_without_writing_audit_log(teacher_context) -> None:
    from server.application.models import OperationLog

    client, database, _ = teacher_context
    question_id = seed_question(
        database,
        "export-unsafe-image-question",
        attachments=[{"kind": "image", "path": "/etc/passwd", "alt": "不可信"}],
    )
    paper = create_manual_paper(client, [question_id])

    response = client.get(
        f"/api/teacher/papers/{paper['id']}/export?variant=questions&format=docx",
        headers=csrf(),
    )

    assert response.status_code == 422
    assert response.json()["code"] == "PAPER_EXPORT_ASSET_INVALID"
    with database.session() as session:
        logs = session.scalars(
            select(OperationLog).where(
                OperationLog.action == "paper.export",
                OperationLog.target_id == paper["id"],
            )
        ).all()
    assert logs == []
