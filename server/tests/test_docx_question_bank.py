from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from PIL import Image
import pytest

from server.application.importers.docx_question_bank import parse_question_bank
from server.application.importers.report import ReportValidationError, validate_generated_files


REAL_SOURCE = Path(__file__).resolve().parents[2] / ".source-question-banks" / "soil-mechanics"


def build_docx_fixture(tmp_path: Path) -> Path:
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    image_path = tmp_path / "seepage.png"
    Image.new("RGB", (12, 8), color=(12, 90, 160)).save(image_path)

    document = Document()
    document.add_heading("第二章 土的渗流", level=1)
    stem = document.add_paragraph("1.(单选题)达西定律适用于哪种流态？")
    stem.add_run().add_picture(str(image_path), width=Inches(0.2))
    document.add_paragraph("A.层流")
    document.add_paragraph("B.紊流")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "水力坡降 i"
    table.cell(0, 1).text = "0.5"
    table.cell(1, 0).text = "渗透系数 k"
    table.cell(1, 1).text = "1.0e-5 m/s"
    document.add_paragraph("正确答案：A")

    source = source_dir / "synthetic.docx"
    document.save(source)
    return source


def load_manifest(output_dir: Path) -> dict:
    return json.loads((output_dir / "manifest.json").read_text(encoding="utf-8"))


def append_omml(paragraph, text: str) -> None:
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    value = OxmlElement("m:t")
    value.text = text
    run.append(value)
    math.append(run)
    paragraph._p.append(math)


def append_superscript_omml(paragraph, base: str, superscript: str) -> None:
    math = OxmlElement("m:oMath")
    superscript_node = OxmlElement("m:sSup")
    base_node = OxmlElement("m:e")
    base_run = OxmlElement("m:r")
    base_text = OxmlElement("m:t")
    base_text.text = base
    base_run.append(base_text)
    base_node.append(base_run)
    exponent_node = OxmlElement("m:sup")
    exponent_run = OxmlElement("m:r")
    exponent_text = OxmlElement("m:t")
    exponent_text.text = superscript
    exponent_run.append(exponent_text)
    exponent_node.append(exponent_run)
    superscript_node.extend((base_node, exponent_node))
    math.append(superscript_node)
    paragraph._p.append(math)


def all_preserved_occurrences(manifest: dict, report: dict) -> list[dict]:
    return [*manifest["questions"], *report["reviewItems"], *report["duplicateOccurrences"]]


def test_parser_preserves_question_media_table_and_source_location(tmp_path: Path) -> None:
    source = build_docx_fixture(tmp_path)
    output_dir = tmp_path / "out"

    report = parse_question_bank(source.parent, output_dir)

    manifest = load_manifest(output_dir)
    item = manifest["questions"][0]
    assert item["subjectId"] == "soil-mechanics"
    assert 1 <= len(item["knowledgePointIds"]) <= 3
    assert item["correctAnswer"] == "A"
    assert item["attachments"][0]["kind"] == "image"
    assert item["attachments"][1] == {
        "kind": "table",
        "inlineOrdinal": 1,
        "placeholder": "[[attachment:1]]",
        "rows": [["水力坡降 i", "0.5"], ["渗透系数 k", "1.0e-5 m/s"]],
        "sourcePosition": {"blockIndex": 5, "tableIndex": 1},
    }
    assert item["sourceMetadata"]["file"] == source.name
    assert item["sourceMetadata"]["position"] == {
        "blockIndex": 2,
        "paragraphIndex": 2,
    }
    assert item["sourceMetadata"]["originalTypeLabel"] == "单选题"
    assert item["sourceMetadata"]["sequence"] == 1
    assert [block["sourcePosition"]["blockIndex"] for block in item["sourceBlocks"]] == [2, 3, 4, 5, 6]
    assert [block["role"] for block in item["sourceBlocks"]] == ["stem", "option", "option", "content", "answer"]
    assert report.unrecognized == 0


def test_parser_preserves_omml_source_and_uses_manual_calculation_contract(tmp_path: Path) -> None:
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    document = Document()
    document.add_heading("第三章 土中应力", level=1)
    stem = document.add_paragraph("1.(计算题)计算地基中的竖向附加应力 ")
    append_omml(stem, "σz=25kPa")
    document.add_paragraph("答案：σz=25kPa")
    source = source_dir / "math.docx"
    document.save(source)

    parse_question_bank(source_dir, tmp_path / "out")

    item = load_manifest(tmp_path / "out")["questions"][0]
    assert item["questionType"] == "计算题"
    assert item["correctAnswer"] is None
    assert item["gradingMode"] == "manual"
    assert item["sourceMetadata"]["sourceAnswer"] == "σz=25kPa"
    formula = item["attachments"][0]
    assert formula["kind"] == "formula"
    assert formula["ommlText"] == "σz=25kPa"
    assert "<m:oMath" in formula["ommlSource"]
    assert formula["sourcePosition"] == {"blockIndex": 2, "paragraphIndex": 2}
    assert formula["inlineOrdinal"] == 1
    assert formula["placeholder"] == "[[attachment:1]]"


def test_uncertain_type_answer_and_mapping_are_reviewed_not_activated(tmp_path: Path) -> None:
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    document = Document()
    document.add_heading("第二章 土的渗流", level=1)
    document.add_paragraph("1.(判断题)达西定律适用于层流。")
    document.add_paragraph("答案：不确定")
    document.add_paragraph("2.(名词解释题)解释达西定律。")
    document.add_paragraph("答案：渗流量与水力坡降成正比。")
    document.add_heading("第九章 未编目内容", level=1)
    document.add_paragraph("3.(判断题)这是一道无法归类的题。")
    document.add_paragraph("答案：正确")
    source = source_dir / "review.docx"
    document.save(source)

    report = parse_question_bank(source_dir, tmp_path / "out")

    manifest = load_manifest(tmp_path / "out")
    report_data = json.loads((tmp_path / "out" / "import-report.json").read_text(encoding="utf-8"))
    assert manifest["questions"] == []
    assert report_data["status"] == "DONE_WITH_CONCERNS"
    assert report_data["sourceQuestions"] == 3
    assert report_data["reviewQuestions"] == 3
    assert report.unrecognized == 0
    reason_codes = {reason["code"] for item in report_data["reviewItems"] for reason in item["reasons"]}
    assert reason_codes == {"INVALID_BOOLEAN_ANSWER", "UNSUPPORTED_TYPE", "UNMAPPED_KNOWLEDGE_POINT"}
    assert [item["sourceMetadata"]["position"]["blockIndex"] for item in report_data["reviewItems"]] == [2, 4, 7]
    assert all(item["sourceMetadata"]["file"] == source.name for item in report_data["reviewItems"])


def test_curated_catalog_has_about_fifty_meaningful_chapter_points(tmp_path: Path) -> None:
    source = build_docx_fixture(tmp_path)
    parse_question_bank(source.parent, tmp_path / "out")

    catalog = json.loads((tmp_path / "out" / "knowledge-points.json").read_text(encoding="utf-8"))
    assert 45 <= len(catalog) <= 55
    assert len({item["id"] for item in catalog}) == len(catalog)
    assert all(item["subjectId"] == "soil-mechanics" for item in catalog)
    assert all(item["status"] == "active" and item["description"] for item in catalog)
    assert [item["sortOrder"] for item in catalog] == sorted(item["sortOrder"] for item in catalog)
    assert {item["chapter"] for item in catalog} == {
        "绪论",
        "第一章 土的性质及工程分类",
        "第二章 土的渗流",
        "第三章 土中应力",
        "第四章 土的压缩性与地基沉降",
        "第五章 土的抗剪强度",
        "第六章 土压力",
        "第七章 土坡稳定性",
        "第八章 地基承载力",
    }


def test_duplicate_content_and_assets_are_deterministic_across_runs(tmp_path: Path) -> None:
    source = build_docx_fixture(tmp_path)
    duplicate = source.with_name("duplicate.docx")
    shutil.copyfile(source, duplicate)

    first_output = tmp_path / "first"
    second_output = tmp_path / "second"
    first_assets = tmp_path / "first-assets"
    second_assets = tmp_path / "second-assets"
    parse_question_bank(source.parent, first_output, first_assets)
    parse_question_bank(source.parent, second_output, second_assets)

    first_manifest = load_manifest(first_output)
    first_report = json.loads((first_output / "import-report.json").read_text(encoding="utf-8"))
    assert len(first_manifest["questions"]) == 1
    assert first_report["sourceQuestions"] == 2
    assert first_report["deduplicatedQuestions"] == 1
    assert first_report["duplicateCount"] == 1
    assert first_report["exactDuplicateCount"] == 1
    assert first_report["conflictedDuplicateCount"] == 0
    assert len(first_report["duplicateOccurrences"]) == 1
    assert first_report["duplicateOccurrences"][0]["classification"] == "exact"
    assert first_report["duplicateOccurrences"][0]["attachments"]
    assert first_report["duplicateOccurrences"][0]["sourceBlocks"]
    assert [item["file"] for item in first_manifest["questions"][0]["sourceMetadata"]["duplicateSources"]] == [
        "synthetic.docx"
    ]
    for filename in ("knowledge-points.json", "manifest.json", "import-report.json"):
        assert (first_output / filename).read_bytes() == (second_output / filename).read_bytes()
    first_asset_files = sorted(path.name for path in first_assets.iterdir())
    second_asset_files = sorted(path.name for path in second_assets.iterdir())
    assert first_asset_files == second_asset_files
    assert all((first_assets / name).read_bytes() == (second_assets / name).read_bytes() for name in first_asset_files)


def test_semantic_dedupe_distinguishes_table_image_and_formula_identity(tmp_path: Path) -> None:
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    red_image = tmp_path / "red.png"
    blue_image = tmp_path / "blue.png"
    Image.new("RGB", (8, 8), color=(220, 30, 30)).save(red_image)
    Image.new("RGB", (8, 8), color=(20, 80, 220)).save(blue_image)
    document = Document()
    document.add_heading("第四章 土的压缩性与地基沉降", level=1)

    for number, value, answer in ((1, "0.10", "0.44"), (2, "0.20", "0.34")):
        document.add_paragraph(f"{number}.(填空题)某土的固结试验结果如下表所示，其压缩系数为____。")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "压力"
        table.cell(0, 1).text = value
        document.add_paragraph(f"答案：{answer}")

    for number, image_path in ((3, red_image), (4, blue_image)):
        paragraph = document.add_paragraph(f"{number}.(计算题)根据附图计算结果。")
        paragraph.add_run().add_picture(str(image_path), width=Inches(0.15))
        document.add_paragraph("答案：1")

    plain_formula = document.add_paragraph("5.(计算题)根据公式x2计算结果。")
    append_omml(plain_formula, "x2")
    document.add_paragraph("答案：1")
    superscript_formula = document.add_paragraph("6.(计算题)根据公式x2计算结果。")
    append_superscript_omml(superscript_formula, "x", "2")
    document.add_paragraph("答案：1")
    document.save(source_dir / "semantic.docx")

    parse_question_bank(source_dir, tmp_path / "out")

    manifest = load_manifest(tmp_path / "out")
    report = json.loads((tmp_path / "out" / "import-report.json").read_text(encoding="utf-8"))
    assert len(manifest["questions"]) == 6
    assert len({item["contentFingerprint"] for item in manifest["questions"]}) == 6
    assert report["exactDuplicateCount"] == 0
    assert report["conflictedDuplicateCount"] == 0
    assert [item["attachments"][0]["rows"][0][1] for item in manifest["questions"][:2]] == ["0.10", "0.20"]
    assert manifest["questions"][2]["attachments"][0]["sha256"] != manifest["questions"][3]["attachments"][0]["sha256"]
    assert manifest["questions"][4]["attachments"][0]["ommlSha256"] != manifest["questions"][5]["attachments"][0]["ommlSha256"]


def test_answer_marker_attachments_keep_position_and_inline_order(tmp_path: Path) -> None:
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    answer_image = tmp_path / "answer.png"
    Image.new("RGB", (10, 6), color=(40, 140, 80)).save(answer_image)
    document = Document()
    document.add_heading("第三章 土中应力", level=1)
    document.add_paragraph("1.(计算题)计算竖向附加应力。")
    answer = document.add_paragraph("答案：")
    answer.add_run().add_picture(str(answer_image), width=Inches(0.15))
    append_omml(answer, "σz=25kPa")
    document.save(source_dir / "answer-media.docx")

    parse_question_bank(source_dir, tmp_path / "out")

    item = load_manifest(tmp_path / "out")["questions"][0]
    assert [attachment["kind"] for attachment in item["attachments"]] == ["image", "formula"]
    assert all(attachment["sourcePosition"] == item["sourceMetadata"]["answerPosition"] for attachment in item["attachments"])
    answer_block = item["sourceBlocks"][-1]
    assert answer_block["role"] == "answer"
    assert answer_block["sourcePosition"] == {"blockIndex": 3, "paragraphIndex": 3}
    assert answer_block["textWithPlaceholders"] == "答案：[[attachment:1]][[attachment:2]]"
    assert answer_block["inlineContent"] == [
        {"kind": "text", "text": "答案："},
        {"attachmentOrdinal": 1, "kind": "attachment", "placeholder": "[[attachment:1]]"},
        {"attachmentOrdinal": 2, "kind": "attachment", "placeholder": "[[attachment:2]]"},
    ]


def test_source_blocks_reconstruct_paragraph_and_inline_media_formula_order(tmp_path: Path) -> None:
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    image_path = tmp_path / "inline.png"
    Image.new("RGB", (7, 7), color=(80, 40, 160)).save(image_path)
    document = Document()
    document.add_heading("第三章 土中应力", level=1)
    stem = document.add_paragraph("1.(计算题)前")
    stem.add_run().add_picture(str(image_path), width=Inches(0.12))
    stem.add_run("中")
    append_omml(stem, "σz")
    stem.add_run("后")
    document.add_paragraph("第二段条件。")
    document.add_paragraph("答案：25")
    document.save(source_dir / "ordered.docx")

    parse_question_bank(source_dir, tmp_path / "out")

    item = load_manifest(tmp_path / "out")["questions"][0]
    assert [block["text"] for block in item["sourceBlocks"]] == [
        "1.(计算题)前中σz后",
        "第二段条件。",
        "答案:25",
    ]
    assert item["sourceBlocks"][0]["textWithPlaceholders"] == "1.(计算题)前[[attachment:1]]中[[attachment:2]]后"
    assert item["sourceBlocks"][1]["role"] == "stemContinuation"
    assert [token["kind"] for token in item["sourceBlocks"][0]["inlineContent"]] == [
        "text", "attachment", "text", "attachment", "text",
    ]


def test_active_review_duplicate_preserves_review_payload_and_reasons(tmp_path: Path) -> None:
    source_dir = tmp_path / "source"
    source_dir.mkdir()
    active = Document()
    active.add_heading("第四章 土的压缩性与地基沉降", level=1)
    active.add_paragraph("1.(简答题)地震、风载会不会引起地基沉降？为什么？")
    active.add_paragraph("答案：")
    active.save(source_dir / "active.docx")
    review = Document()
    review.add_heading("第四章 土的压缩性与地基沉降", level=1)
    review.add_paragraph("1.(名词解释题)地震、风载会不会引起地基沉降？为什么？")
    review.add_paragraph("答案：会引起附加沉降，因为动力荷载改变土体应力与结构。")
    review.save(source_dir / "review.docx")

    parse_question_bank(source_dir, tmp_path / "out")

    manifest = load_manifest(tmp_path / "out")
    report = json.loads((tmp_path / "out" / "import-report.json").read_text(encoding="utf-8"))
    assert len(manifest["questions"]) == 1
    assert manifest["questions"][0]["sourceMetadata"]["sourceAnswer"] == ""
    assert report["reviewQuestions"] == 1
    assert report["reviewGatedOccurrenceCount"] == 1
    assert report["conflictedDuplicateCount"] == 1
    assert report["exactDuplicateCount"] == 0
    review_item = report["reviewItems"][0]
    assert review_item["sourceAnswer"].startswith("会引起附加沉降")
    assert {reason["code"] for reason in review_item["reasons"]} >= {"UNSUPPORTED_TYPE"}
    assert review_item["duplicateClassification"] == "conflicted"
    assert review_item["sourceBlocks"][-1]["role"] == "answer"


def test_archive_provenance_is_unset_for_filename_only_match(tmp_path: Path) -> None:
    source = build_docx_fixture(tmp_path)
    for name in (
        "1-随堂检测题库.docx",
        "2-章节测验.docx",
        "3-其他题库-1.docx",
        "4-其他题库-2.docx",
    ):
        shutil.copyfile(source, source.parent / name)
    source.unlink()
    (source.parent / "说明.txt").write_text("untrusted synthetic note", encoding="utf-8")

    parse_question_bank(source.parent, tmp_path / "out")

    manifest = load_manifest(tmp_path / "out")
    report = json.loads((tmp_path / "out" / "import-report.json").read_text(encoding="utf-8"))
    assert manifest["sourceArchiveSha256"] is None
    assert report["sourceArchiveSha256"] is None
    assert report["sourceArchiveVerified"] is False


@pytest.mark.skipif(not REAL_SOURCE.exists(), reason="ignored soil-mechanics source corpus is not available")
def test_real_corpus_covers_all_sources_and_accounts_for_every_candidate(tmp_path: Path) -> None:
    assets_dir = tmp_path / "assets"
    parse_question_bank(REAL_SOURCE, tmp_path / "out", assets_dir)

    report = json.loads((tmp_path / "out" / "import-report.json").read_text(encoding="utf-8"))
    manifest = load_manifest(tmp_path / "out")
    assert [item["file"] for item in report["sourceFiles"]] == [
        "1-随堂检测题库.docx",
        "2-章节测验.docx",
        "3-其他题库-1.docx",
        "4-其他题库-2.docx",
    ]
    assert report["sourceFileCount"] == 4
    assert report["answerMarkerCount"] >= 1000
    assert report["sourceQuestions"] >= 1000
    assert report["imageCount"] >= 300
    assert report["tableCount"] >= 8
    assert report["sourceQuestions"] == report["deduplicatedQuestions"] + report["duplicateCount"]
    assert report["deduplicatedQuestions"] == report["activeQuestions"] + report["reviewQuestions"]
    assert report["duplicateCount"] == report["exactDuplicateCount"]
    assert report["reviewQuestions"] == report["reviewGatedOccurrenceCount"]
    assert report["imageAssetCount"] == len(list(assets_dir.iterdir()))
    assert all(item["sourceMetadata"]["file"] in {source["file"] for source in report["sourceFiles"]} for item in report["reviewItems"])
    assert report["sourceArchiveVerified"] is True

    sequence_items = [
        item for item in manifest["questions"]
        if item["sourceMetadata"]["file"] == "3-其他题库-1.docx"
        and item["sourceMetadata"]["sequence"] in {39, 40, 41, 42}
    ]
    assert {item["sourceMetadata"]["sequence"] for item in sequence_items} == {39, 40, 41, 42}
    assert len({tuple(tuple(row) for row in item["attachments"][0]["rows"]) for item in sequence_items}) == 4

    accounting = report["attachmentAccounting"]
    assert accounting["sourceOccurrenceCount"] == report["imageCount"] + report["tableCount"] + report["ommlCount"]
    assert accounting["accountedOccurrenceCount"] == accounting["sourceOccurrenceCount"]
    assert accounting["unaccountedOccurrenceCount"] == 0
    preserved = all_preserved_occurrences(manifest, report)
    referenced_assets = {
        Path(attachment["src"]).name
        for item in preserved
        for attachment in item["attachments"]
        if attachment["kind"] == "image"
    }
    referenced_assets.update(
        Path(record["attachment"]["src"]).name
        for record in report["attachmentLossRecords"]
        if record["attachment"]["kind"] == "image"
    )
    assert referenced_assets == {path.name for path in assets_dir.iterdir()}

    sequence_85 = next(
        item for item in preserved
        if item["sourceMetadata"]["file"] == "3-其他题库-1.docx"
        and item["sourceMetadata"]["sequence"] == 85
    )
    assert any(
        attachment["kind"] == "image"
        and attachment["sourcePosition"] == sequence_85["sourceMetadata"]["answerPosition"]
        for attachment in sequence_85["attachments"]
    )
    assert all(
        item["sourceBlocks"]
        and item["sourceBlocks"][0]["sourcePosition"] == item["sourceMetadata"]["position"]
        for item in preserved
    )


def test_report_validator_enforces_strict_active_question_contract(tmp_path: Path) -> None:
    source = build_docx_fixture(tmp_path)
    output_dir = tmp_path / "out"
    parse_question_bank(source.parent, output_dir)
    manifest_path = output_dir / "manifest.json"
    report_path = output_dir / "import-report.json"

    summary = validate_generated_files(manifest_path, report_path)

    assert summary == {"activeQuestions": 1, "reviewQuestions": 0, "status": "DONE"}
    manifest = load_manifest(output_dir)
    manifest["questions"][0]["correctAnswer"] = ["A"]
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
    with pytest.raises(ReportValidationError, match="SINGLE_CHOICE_ANSWER"):
        validate_generated_files(manifest_path, report_path)
