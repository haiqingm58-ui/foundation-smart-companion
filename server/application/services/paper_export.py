from __future__ import annotations

import base64
import binascii
import hashlib
import io
import math
import os
import re
from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import Any, BinaryIO, Literal

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image as PillowImage
from PIL import UnidentifiedImageError
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ExportVariant = Literal["questions", "answer-sheet", "answers"]
ExportFormat = Literal["docx", "pdf"]
ContentSide = Literal["question", "answer"]

DOCX_FONT_NAME = "Arial Unicode MS"
PDF_FONT_NAME = "FoundationPaperUnicode"
APP_ASSET_PREFIX = "/foundation-smart-companion/question-assets/soil-mechanics/"
MAX_IMAGE_BYTES = 8 * 1024 * 1024
MAX_EMBEDDED_IMAGE_BYTES = 20 * 1024 * 1024
MAX_IMAGE_PIXELS = 25_000_000
MAX_TABLE_ROW_LINES = 24
ALLOWED_IMAGE_FORMATS = {
    "png": "PNG",
    "jpg": "JPEG",
    "jpeg": "JPEG",
    "gif": "GIF",
    "webp": "WEBP",
}
ANSWER_ROLES = {"answer", "postanswer", "explanation", "analysis", "solution"}


class PaperExportError(ValueError):
    pass


@dataclass(frozen=True)
class ExportOptions:
    include_student_fields: bool = True
    asset_root: Path | None = None


@dataclass(frozen=True)
class ContentUnit:
    kind: Literal["text", "formula", "table", "image"]
    value: Any


@dataclass(frozen=True)
class SourceBlockContent:
    role: str
    units: tuple[ContentUnit, ...]


def render_paper(
    paper: Any,
    variant: ExportVariant,
    format: ExportFormat,
    options: ExportOptions,
) -> bytes:
    output = io.BytesIO()
    render_paper_to_file(paper, variant, format, options, output)
    return output.getvalue()


def render_paper_to_file(
    paper: Any,
    variant: ExportVariant,
    format: ExportFormat,
    options: ExportOptions,
    output: BinaryIO,
) -> None:
    if variant not in {"questions", "answer-sheet", "answers"}:
        raise ValueError(f"unsupported paper export variant: {variant}")
    if format == "docx":
        _render_docx(paper, variant, options, output)
        return
    if format == "pdf":
        _render_pdf(paper, variant, options, output)
        return
    raise ValueError(f"unsupported paper export format: {format}")


def _value(record: Any, *names: str, default: Any = None) -> Any:
    for name in names:
        if isinstance(record, dict) and name in record:
            return record[name]
        if hasattr(record, name):
            return getattr(record, name)
    return default


def _paper_questions(paper: Any) -> list[Any]:
    rows = _value(paper, "paper_questions", "paperQuestions", default=[])
    return sorted(rows or [], key=lambda item: _value(item, "sequence", default=0))


def _question(row: Any) -> Any:
    return _value(row, "question", default=row)


def _question_value(row: Any, *names: str, default: Any = None) -> Any:
    return _value(_question(row), *names, default=default)


def _source_metadata(row: Any) -> dict[str, Any]:
    value = _question_value(row, "source_metadata", "sourceMetadata", default={})
    return value if isinstance(value, dict) else {}


def _source_blocks(row: Any) -> list[dict[str, Any]]:
    direct = _question_value(row, "source_blocks", "sourceBlocks", default=None)
    if isinstance(direct, list):
        return [block for block in direct if isinstance(block, dict)]
    stored = _value(_source_metadata(row), "sourceBlocks", "source_blocks", default=[])
    return [block for block in stored if isinstance(block, dict)] if isinstance(stored, list) else []


def _paper_title(paper: Any, variant: ExportVariant) -> str:
    title = str(_value(paper, "title", default="试卷"))
    return f"{title}答题卡" if variant == "answer-sheet" else title


def _subject_name(paper: Any) -> str:
    subject = _value(paper, "subject", default=None)
    return str(_value(subject, "name", "title", default="") or "")


def _format_number(value: Any) -> str:
    if isinstance(value, bool):
        return "1" if value else "0"
    if isinstance(value, (int, float)):
        return f"{value:g}"
    return str(value)


def _metadata_text(paper: Any) -> str:
    duration = _value(paper, "duration_minutes", "durationMinutes", default=None)
    duration_text = "考试时间：未限定" if duration is None else f"考试时间：{_format_number(duration)} 分钟"
    total = _value(paper, "total_points", "totalPoints", default=0)
    return "  ".join(
        item
        for item in (
            f"课程：{_subject_name(paper)}" if _subject_name(paper) else "",
            duration_text,
            f"满分：{_format_number(total)} 分",
        )
        if item
    )


def _format_answer_value(value: Any) -> str:
    if isinstance(value, bool):
        return "正确" if value else "错误"
    if isinstance(value, (list, tuple, set)):
        return "、".join(_format_answer_value(item) for item in value)
    if value is None:
        return "未提供"
    return str(value)


def _answer_text(row: Any) -> str:
    value = _question_value(row, "correct_answer", "correctAnswer")
    question_type = str(_question_value(row, "question_type", "questionType", default=""))
    if value is None:
        return "未提供"
    if question_type == "判断题" or isinstance(value, bool):
        if isinstance(value, str):
            normalized = value.strip().casefold()
            if normalized in {"true", "1", "yes", "正确", "对", "是"}:
                return "正确"
            if normalized in {"false", "0", "no", "错误", "错", "否"}:
                return "错误"
        return _format_answer_value(value)
    if isinstance(value, (list, tuple, set)):
        joined = "、".join(_format_answer_value(item) for item in value)
        return f"可接受答案：{joined}" if question_type == "填空题" else joined
    if not isinstance(value, dict):
        return _format_answer_value(value)

    field_labels = {
        "aliases": "同义答案",
        "acceptedAnswers": "可接受答案",
        "tolerance": "容差",
        "precision": "精度",
        "method": "方法",
        "steps": "步骤",
    }
    parts: list[str] = []
    if "value" in value:
        primary = _format_answer_value(value["value"])
        unit = value.get("unit")
        if unit not in {None, ""}:
            primary = f"{primary} {_format_answer_value(unit)}"
        parts.append(primary)
    ordered_keys = [key for key in field_labels if key in value]
    ordered_keys.extend(sorted(key for key in value if key not in {"value", "unit", *field_labels}))
    if "value" not in value and "unit" in value:
        ordered_keys.insert(0, "unit")
    for key in ordered_keys:
        label = field_labels.get(key, "单位" if key == "unit" else str(key))
        parts.append(f"{label}：{_format_answer_value(value[key])}")
    return "；".join(parts) if parts else "未提供"


def _block_index(position: Any) -> int | None:
    if not isinstance(position, dict):
        return None
    value = position.get("blockIndex")
    return value if isinstance(value, int) and not isinstance(value, bool) else None


def _position_key(position: Any, ordinal: Any = 0) -> tuple[int, int, int, int]:
    if not isinstance(position, dict):
        return (10**9, 10**9, 10**9, int(ordinal or 0))
    return (
        int(position.get("blockIndex", 10**9)),
        int(position.get("paragraphIndex", 10**9)),
        int(position.get("tableIndex", 10**9)),
        int(ordinal or 0),
    )


def _answer_boundary(row: Any) -> int | None:
    return _block_index(_source_metadata(row).get("answerPosition"))


def _block_side(block: dict[str, Any], answer_boundary: int | None) -> ContentSide:
    role = str(block.get("role", "")).casefold()
    if role in ANSWER_ROLES:
        return "answer"
    block_index = _block_index(block.get("sourcePosition"))
    if answer_boundary is not None and block_index is not None and block_index >= answer_boundary:
        return "answer"
    return "question"


def _attachment_side(
    attachment: Any,
    blocks: list[dict[str, Any]],
    answer_boundary: int | None,
) -> ContentSide:
    position = _value(attachment, "sourcePosition", "source_position", default=None)
    for block in blocks:
        if block.get("sourcePosition") == position:
            return _block_side(block, answer_boundary)
    index = _block_index(position)
    if answer_boundary is not None and index is not None and index >= answer_boundary:
        return "answer"
    return "question"


def _attachment_unit(attachment: Any) -> ContentUnit | None:
    kind = _value(attachment, "kind", "type")
    if kind in {"image", "table", "formula"}:
        return ContentUnit(kind, attachment)
    return None


def _block_contents(row: Any, side: ContentSide) -> list[SourceBlockContent]:
    blocks = _source_blocks(row)
    if not blocks:
        return []
    attachments = list(_question_value(row, "attachments", default=[]) or [])
    answer_boundary = _answer_boundary(row)
    consumed: set[int] = set()
    result: list[SourceBlockContent] = []
    for block in blocks:
        if _block_side(block, answer_boundary) != side:
            continue
        position = block.get("sourcePosition")
        matching = [
            (index, attachment)
            for index, attachment in enumerate(attachments)
            if index not in consumed
            and _value(attachment, "sourcePosition", "source_position", default=None) == position
            and _attachment_side(attachment, blocks, answer_boundary) == side
        ]
        units: list[ContentUnit] = []
        inline_content = block.get("inlineContent")
        if isinstance(inline_content, list):
            for token in inline_content:
                if not isinstance(token, dict):
                    continue
                if token.get("kind") == "text" and token.get("text"):
                    units.append(ContentUnit("text", str(token["text"])))
                    continue
                if token.get("kind") != "attachment":
                    continue
                ordinal = token.get("attachmentOrdinal")
                match = next(
                    (
                        (index, attachment)
                        for index, attachment in matching
                        if index not in consumed
                        and _value(attachment, "inlineOrdinal", "inline_ordinal", default=None) == ordinal
                    ),
                    None,
                )
                if match is None:
                    continue
                consumed.add(match[0])
                unit = _attachment_unit(match[1])
                if unit:
                    units.append(unit)
        elif block.get("text"):
            units.append(ContentUnit("text", str(block["text"])))
        for index, attachment in matching:
            if index in consumed:
                continue
            consumed.add(index)
            unit = _attachment_unit(attachment)
            if unit:
                units.append(unit)
        result.append(SourceBlockContent(str(block.get("role", "content")), tuple(units)))

    remaining = [
        attachment
        for index, attachment in enumerate(attachments)
        if index not in consumed and _attachment_side(attachment, blocks, answer_boundary) == side
    ]
    remaining.sort(
        key=lambda attachment: _position_key(
            _value(attachment, "sourcePosition", "source_position", default=None),
            _value(attachment, "inlineOrdinal", "inline_ordinal", default=0),
        )
    )
    units = tuple(unit for attachment in remaining if (unit := _attachment_unit(attachment)))
    if units:
        result.append(SourceBlockContent("content", units))
    return result


def _formula_text(value: Any) -> str:
    if isinstance(value, str):
        return value
    formula = _value(value, "ommlText", "omml_text", "latex", "text", "value", default="")
    return str(formula or "")


def _fallback_units(row: Any, side: ContentSide) -> list[ContentUnit]:
    if side == "answer":
        attachments = list(_question_value(row, "attachments", default=[]) or [])
        boundary = _answer_boundary(row)
        return [
            unit
            for attachment in attachments
            if _attachment_side(attachment, [], boundary) == "answer"
            if (unit := _attachment_unit(attachment))
        ]

    metadata = _source_metadata(row)
    units: list[ContentUnit] = []
    for formula in _value(metadata, "formulas", default=[]) or []:
        units.append(ContentUnit("formula", formula))
    for table_rows in _value(metadata, "tables", default=[]) or []:
        units.append(ContentUnit("table", {"kind": "table", "rows": table_rows}))
    boundary = _answer_boundary(row)
    for attachment in _question_value(row, "attachments", default=[]) or []:
        if _attachment_side(attachment, [], boundary) != "question":
            continue
        unit = _attachment_unit(attachment)
        if unit:
            units.append(unit)
    return units


def _decode_data_uri(source: str) -> tuple[bytes, str]:
    match = re.fullmatch(r"data:image/(png|jpeg|gif|webp);base64,([A-Za-z0-9+/=]+)", source)
    if match is None:
        raise PaperExportError("图片数据 URI 类型或编码无效")
    encoded = match.group(2)
    if len(encoded) > math.ceil(MAX_IMAGE_BYTES / 3) * 4 + 4:
        raise PaperExportError("图片大小超过导出限制")
    try:
        data = base64.b64decode(encoded, validate=True)
    except (binascii.Error, ValueError) as error:
        raise PaperExportError("图片数据 URI 编码无效") from error
    return data, ALLOWED_IMAGE_FORMATS[match.group(1)]


def _read_app_asset(source: str, attachment: Any, options: ExportOptions) -> tuple[bytes, str]:
    pattern = re.escape(APP_ASSET_PREFIX) + r"([0-9a-f]{64})\.(png|jpg|jpeg|gif|webp)"
    match = re.fullmatch(pattern, source)
    if match is None:
        raise PaperExportError("图片 URL 不在允许的题库资源目录内")
    expected_digest, suffix = match.groups()
    declared_digest = _value(attachment, "sha256", default=None)
    if declared_digest != expected_digest:
        raise PaperExportError("图片 SHA-256 与资源文件名不一致")
    if options.asset_root is None:
        raise PaperExportError("图片资源目录未配置")
    try:
        root = Path(options.asset_root).resolve(strict=True)
    except OSError as error:
        raise PaperExportError("图片资源目录不存在") from error
    relative = Path("question-assets") / "soil-mechanics" / f"{expected_digest}.{suffix}"
    candidate = (root / relative).resolve()
    if not candidate.is_relative_to(root):
        raise PaperExportError("图片资源越出允许目录")
    if not candidate.is_file():
        raise PaperExportError("图片资源不存在")
    try:
        with candidate.open("rb") as handle:
            data = handle.read(MAX_IMAGE_BYTES + 1)
    except OSError as error:
        raise PaperExportError("图片资源无法读取") from error
    if len(data) > MAX_IMAGE_BYTES:
        raise PaperExportError("图片大小超过导出限制")
    if hashlib.sha256(data).hexdigest() != expected_digest:
        raise PaperExportError("图片内容 SHA-256 校验失败")
    return data, ALLOWED_IMAGE_FORMATS[suffix]


def _validated_png(data: bytes, expected_format: str, attachment: Any) -> bytes:
    if len(data) > MAX_IMAGE_BYTES:
        raise PaperExportError("图片大小超过导出限制")
    declared_digest = _value(attachment, "sha256", default=None)
    if declared_digest is not None:
        if not isinstance(declared_digest, str) or re.fullmatch(r"[0-9a-f]{64}", declared_digest) is None:
            raise PaperExportError("图片 SHA-256 格式无效")
        if hashlib.sha256(data).hexdigest() != declared_digest:
            raise PaperExportError("图片内容 SHA-256 校验失败")
    try:
        with PillowImage.open(io.BytesIO(data)) as probe:
            actual_format = str(probe.format or "").upper()
            width, height = probe.size
            probe.verify()
        if actual_format != expected_format:
            raise PaperExportError("图片声明类型与实际类型不一致")
        if width <= 0 or height <= 0 or width * height > MAX_IMAGE_PIXELS:
            raise PaperExportError("图片尺寸超过导出限制")
        with PillowImage.open(io.BytesIO(data)) as image:
            image.seek(0)
            image.load()
            converted = image.convert("RGBA" if image.mode in {"RGBA", "LA"} or "transparency" in image.info else "RGB")
            output = io.BytesIO()
            converted.save(output, format="PNG", compress_level=9)
    except PaperExportError:
        raise
    except (OSError, ValueError, UnidentifiedImageError, PillowImage.DecompressionBombError) as error:
        raise PaperExportError("图片内容无法安全解码") from error
    normalized = output.getvalue()
    if len(normalized) > MAX_EMBEDDED_IMAGE_BYTES:
        raise PaperExportError("图片转换后大小超过导出限制")
    return normalized


def _image_bytes(attachment: Any, options: ExportOptions) -> bytes:
    if _value(attachment, "kind", "type") != "image":
        raise PaperExportError("图片附件类型无效")
    if isinstance(_value(attachment, "path", default=None), str):
        raise PaperExportError("图片来源不受支持，仅允许数据 URI 或应用题库资源")
    source = _value(attachment, "data", "src", "url", default=None)
    if not isinstance(source, str) or not source:
        raise PaperExportError("图片来源缺失")
    if source.startswith("data:"):
        data, expected_format = _decode_data_uri(source)
    else:
        data, expected_format = _read_app_asset(source, attachment, options)
    return _validated_png(data, expected_format, attachment)


def _image_size_cm(data: bytes, max_width: float, max_height: float) -> tuple[float, float]:
    with PillowImage.open(io.BytesIO(data)) as image:
        width, height = image.size
    scale = min(max_width / width, max_height / height)
    return width * scale, height * scale


def _set_docx_font_slots(r_pr, font_name: str = DOCX_FONT_NAME) -> None:
    r_fonts = r_pr.get_or_add_rFonts()
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{slot}"), font_name)
    r_fonts.set(qn("w:hint"), "eastAsia")
    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    for slot in ("val", "eastAsia", "bidi"):
        language.set(qn(f"w:{slot}"), "zh-CN")


def _set_docx_run_font(run, font_name: str = DOCX_FONT_NAME) -> None:
    run.font.name = font_name
    _set_docx_font_slots(run._element.get_or_add_rPr(), font_name)


def _set_docx_style_font(style, font_name: str = DOCX_FONT_NAME) -> None:
    style.font.name = font_name
    _set_docx_font_slots(style._element.get_or_add_rPr(), font_name)


def _set_docx_language_defaults(document: Document) -> None:
    settings = document.settings.element
    theme_language = settings.find(qn("w:themeFontLang"))
    if theme_language is None:
        theme_language = OxmlElement("w:themeFontLang")
        settings.append(theme_language)
    for slot in ("val", "eastAsia", "bidi"):
        theme_language.set(qn(f"w:{slot}"), "zh-CN")
    default_properties = document.styles.element.xpath(
        "./w:docDefaults/w:rPrDefault/w:rPr"
    )[0]
    _set_docx_font_slots(default_properties)


def _set_cell_text(cell, text: str, *, bold: bool = False, font_size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    _set_docx_run_font(run)


def _prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def _repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _set_table_widths(table, widths: list[Cm]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = width


def _table_rows_for_export(rows: list[list[Any]]) -> list[list[str]]:
    if not rows:
        return []
    column_count = max(1, max((len(row) for row in rows), default=1))
    chars_per_line = max(8, 52 // column_count)
    max_chars = chars_per_line * MAX_TABLE_ROW_LINES
    normalized: list[list[str]] = []
    for row in rows:
        cells = [str(row[index]) if index < len(row) else "" for index in range(column_count)]
        chunks = [
            [value[offset : offset + max_chars] for offset in range(0, len(value), max_chars)] or [""]
            for value in cells
        ]
        chunk_count = max(len(parts) for parts in chunks)
        for chunk_index in range(chunk_count):
            values = [parts[chunk_index] if chunk_index < len(parts) else "" for parts in chunks]
            if chunk_index and cells[0] and len(chunks[0]) == 1:
                values[0] = f"{cells[0]}（续 {chunk_index + 1}/{chunk_count}）"
            normalized.append(values)
    return normalized


def _docx_table(document: Document, rows: list[list[Any]]) -> None:
    if not rows:
        return
    rows = _table_rows_for_export(rows)
    column_count = max(1, max((len(row) for row in rows), default=1))
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    _set_table_widths(table, [Cm(16.5 / column_count)] * column_count)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index in range(column_count):
            _set_cell_text(
                cells[index],
                str(values[index]) if index < len(values) else "",
                bold=row_index == 0,
            )
        if row_index == 0:
            _repeat_table_header(table.rows[-1])
        _prevent_row_split(table.rows[-1])
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)


def _docx_text(
    document: Document,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: float = 10.5,
    keep: bool = False,
) -> None:
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.keep_together = len(text) <= 600
    paragraph.paragraph_format.keep_with_next = keep
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    _set_docx_run_font(run)


def _docx_unit(document: Document, unit: ContentUnit, options: ExportOptions) -> None:
    if unit.kind == "text":
        _docx_text(document, str(unit.value))
        return
    if unit.kind == "formula":
        formula = _formula_text(unit.value)
        if formula:
            _docx_text(document, f"公式：{formula}", italic=True, keep=True)
        return
    if unit.kind == "table":
        _docx_table(document, list(_value(unit.value, "rows", default=[]) or []))
        return
    image = _image_bytes(unit.value, options)
    width, height = _image_size_cm(image, 5.0, 4.8)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    _set_docx_run_font(run)
    run.add_picture(io.BytesIO(image), width=Cm(width), height=Cm(height))
    _docx_text(
        document,
        str(_value(unit.value, "alt", "caption", default="题图") or "题图"),
        italic=True,
        size=9,
    )


def _docx_question(document: Document, row: Any, options: ExportOptions) -> None:
    number = _value(row, "sequence", default=0)
    points = _value(row, "points", default=0)
    text = str(_question_value(row, "text", default=""))
    _docx_text(document, f"{number}. {text}（{_format_number(points)} 分）", bold=True, keep=True)
    source_contents = _block_contents(row, "question")
    question_options = list(_question_value(row, "options", default=[]) or [])
    option_index = 0
    if source_contents:
        for content in source_contents:
            role = content.role.casefold()
            if role == "option":
                if option_index < len(question_options):
                    option = question_options[option_index]
                    _docx_text(document, f"{_value(option, 'label', default='')}. {_value(option, 'text', default='')}")
                    option_index += 1
                units = [unit for unit in content.units if unit.kind != "text"]
            elif role == "stem":
                units = [unit for unit in content.units if unit.kind != "text"]
            else:
                units = list(content.units)
            for unit in units:
                _docx_unit(document, unit, options)
        for option in question_options[option_index:]:
            _docx_text(document, f"{_value(option, 'label', default='')}. {_value(option, 'text', default='')}")
        return
    for option in question_options:
        _docx_text(document, f"{_value(option, 'label', default='')}. {_value(option, 'text', default='')}")
    for unit in _fallback_units(row, "question"):
        _docx_unit(document, unit, options)


def _docx_answer(document: Document, row: Any, options: ExportOptions) -> None:
    number = _value(row, "sequence", default=0)
    _docx_text(document, f"{number}. 正确答案：{_answer_text(row)}", bold=True)
    explanation = _question_value(row, "explanation")
    if explanation:
        _docx_text(document, f"解析：{explanation}")
    contents = _block_contents(row, "answer")
    if contents:
        for content in contents:
            units = (
                [unit for unit in content.units if unit.kind != "text"]
                if content.role.casefold() in ANSWER_ROLES
                else content.units
            )
            for unit in units:
                _docx_unit(document, unit, options)
    else:
        for unit in _fallback_units(row, "answer"):
            _docx_unit(document, unit, options)


def _render_docx(
    paper: Any,
    variant: ExportVariant,
    options: ExportOptions,
    output: BinaryIO,
) -> None:
    document = Document()
    _set_docx_language_defaults(document)
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(1.6 if variant == "answers" else 1.8)
    section.left_margin = section.right_margin = Cm(2.2)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)
    normal = document.styles["Normal"]
    _set_docx_style_font(normal)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15
    title_style = document.styles.add_style("ExamTitle", WD_STYLE_TYPE.PARAGRAPH)
    _set_docx_style_font(title_style)
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.paragraph_format.space_after = Pt(6)
    title = document.add_paragraph(style="ExamTitle")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(_paper_title(paper, variant))
    _set_docx_run_font(title_run)
    if options.include_student_fields:
        info = document.add_table(rows=1, cols=3)
        info.style = "Table Grid"
        _set_table_widths(info, [Cm(5.5), Cm(5.5), Cm(5.5)])
        for cell, text in zip(
            info.rows[0].cells,
            ["姓名：____________", "学号：____________", "班级：____________"],
        ):
            _set_cell_text(cell, text)
        _prevent_row_split(info.rows[0])
    _docx_text(document, _metadata_text(paper), size=10)
    rows = _paper_questions(paper)
    if variant == "answer-sheet":
        for row in rows:
            _docx_text(
                document,
                f"{_value(row, 'sequence', default=0)}. （{_format_number(_value(row, 'points', default=0))} 分）",
                bold=True,
            )
            line_count = 3 if _question_value(row, "question_type", "questionType", default="") in {"简答题", "计算题"} else 1
            for _ in range(line_count):
                _docx_text(document, "________________________________________________________________________")
    else:
        previous_section = None
        for row in rows:
            section_title = str(_value(row, "section_title", "sectionTitle", default=""))
            if section_title and section_title != previous_section:
                _docx_text(document, section_title, bold=True, size=13, keep=True)
                previous_section = section_title
            _docx_question(document, row, options)
    if variant == "answers":
        _docx_text(document, "参考答案", bold=True, size=13, keep=True)
        for row in rows:
            _docx_answer(document, row, options)
    document.save(output)


def _pdf_font_candidates() -> list[tuple[Path, int]]:
    candidates: list[tuple[Path, int]] = []
    configured = os.getenv("FOUNDATION_PDF_FONT_PATH", "").strip()
    if configured:
        candidates.append((Path(configured), int(os.getenv("FOUNDATION_PDF_FONT_SUBFONT_INDEX", "0"))))
    candidates.extend(
        [
            (Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"), 0),
            (Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"), 0),
            (Path("/Library/Fonts/Arial Unicode.ttf"), 0),
            (Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"), 0),
            (Path("/System/Library/Fonts/STHeiti Medium.ttc"), 0),
            (Path("/System/Library/Fonts/Supplemental/Songti.ttc"), 0),
        ]
    )
    return candidates


def _register_pdf_font() -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    if PDF_FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return
    errors: list[str] = []
    for font_path, subfont_index in _pdf_font_candidates():
        if not font_path.is_file():
            continue
        try:
            font = TTFont(PDF_FONT_NAME, str(font_path), subfontIndex=subfont_index)
            if not all(ord(character) in font.face.charToGlyph for character in "土力学试卷"):
                errors.append(f"{font_path}: missing Chinese glyphs")
                continue
            pdfmetrics.registerFont(font)
            return
        except Exception as error:  # ReportLab raises several font-parser exception types.
            errors.append(f"{font_path}: {error}")
    detail = "; ".join(errors) if errors else "no candidate font files"
    raise RuntimeError(f"未找到可嵌入的中文 TrueType 字体，不能安全导出 PDF（{detail}）")


def _pdf_styles() -> dict[str, ParagraphStyle]:
    _register_pdf_font()
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ExamTitle", parent=base["Title"], fontName=PDF_FONT_NAME, fontSize=18,
            leading=24, alignment=TA_CENTER, spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "ExamBody", parent=base["BodyText"], fontName=PDF_FONT_NAME, fontSize=10.5,
            leading=15, spaceAfter=3,
        ),
        "section": ParagraphStyle(
            "ExamSection", parent=base["Heading2"], fontName=PDF_FONT_NAME, fontSize=13,
            leading=18, spaceBefore=7, spaceAfter=4, keepWithNext=True,
        ),
        "question": ParagraphStyle(
            "ExamQuestion", parent=base["BodyText"], fontName=PDF_FONT_NAME, fontSize=10.5,
            leading=15, spaceAfter=3, keepWithNext=True,
        ),
        "question_long": ParagraphStyle(
            "ExamQuestionLong", parent=base["BodyText"], fontName=PDF_FONT_NAME, fontSize=10.5,
            leading=15, spaceAfter=3, keepWithNext=False,
        ),
        "small": ParagraphStyle(
            "ExamSmall", parent=base["BodyText"], fontName=PDF_FONT_NAME, fontSize=9,
            leading=12, textColor=colors.HexColor("#555555"), spaceAfter=3,
        ),
    }


def _pdf_paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(str(text)).replace("\n", "<br/>"), style)


def _pdf_table(rows: list[list[Any]], styles: dict[str, ParagraphStyle]) -> Table | None:
    if not rows:
        return None
    rows = _table_rows_for_export(rows)
    column_count = max(1, max((len(row) for row in rows), default=1))
    data = [
        [
            _pdf_paragraph(str(row[index]) if index < len(row) else "", styles["body"])
            for index in range(column_count)
        ]
        for row in rows
    ]
    table = Table(
        data,
        colWidths=[16.6 * cm / column_count] * column_count,
        repeatRows=1,
        splitByRow=1,
        splitInRow=0,
        hAlign="LEFT",
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#777777")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _pdf_unit(
    story: list[Any],
    unit: ContentUnit,
    options: ExportOptions,
    styles: dict[str, ParagraphStyle],
) -> None:
    if unit.kind == "text":
        if str(unit.value):
            story.append(_pdf_paragraph(str(unit.value), styles["body"]))
        return
    if unit.kind == "formula":
        formula = _formula_text(unit.value)
        if formula:
            story.append(_pdf_paragraph(f"公式：{formula}", styles["small"]))
        return
    if unit.kind == "table":
        table = _pdf_table(list(_value(unit.value, "rows", default=[]) or []), styles)
        if table:
            story.extend([table, Spacer(1, 3)])
        return
    image = _image_bytes(unit.value, options)
    width, height = _image_size_cm(image, 10.0, 5.5)
    picture = PdfImage(io.BytesIO(image), width=width * cm, height=height * cm)
    picture.hAlign = "CENTER"
    story.extend(
        [
            picture,
            _pdf_paragraph(
                str(_value(unit.value, "alt", "caption", default="题图") or "题图"),
                styles["small"],
            ),
        ]
    )


def _pdf_question(
    story: list[Any],
    row: Any,
    options: ExportOptions,
    styles: dict[str, ParagraphStyle],
) -> None:
    number = _value(row, "sequence", default=0)
    points = _value(row, "points", default=0)
    question_text = str(_question_value(row, "text", default=""))
    question_style = styles["question_long"] if len(question_text) > 500 else styles["question"]
    story.append(
        _pdf_paragraph(
            f"{number}. {question_text}（{_format_number(points)} 分）",
            question_style,
        )
    )
    source_contents = _block_contents(row, "question")
    question_options = list(_question_value(row, "options", default=[]) or [])
    option_index = 0
    if source_contents:
        for content in source_contents:
            role = content.role.casefold()
            if role == "option":
                if option_index < len(question_options):
                    option = question_options[option_index]
                    story.append(
                        _pdf_paragraph(
                            f"{_value(option, 'label', default='')}. {_value(option, 'text', default='')}",
                            styles["body"],
                        )
                    )
                    option_index += 1
                units = [unit for unit in content.units if unit.kind != "text"]
            elif role == "stem":
                units = [unit for unit in content.units if unit.kind != "text"]
            else:
                units = list(content.units)
            for unit in units:
                _pdf_unit(story, unit, options, styles)
        for option in question_options[option_index:]:
            story.append(
                _pdf_paragraph(
                    f"{_value(option, 'label', default='')}. {_value(option, 'text', default='')}",
                    styles["body"],
                )
            )
        return
    for option in question_options:
        story.append(
            _pdf_paragraph(
                f"{_value(option, 'label', default='')}. {_value(option, 'text', default='')}",
                styles["body"],
            )
        )
    for unit in _fallback_units(row, "question"):
        _pdf_unit(story, unit, options, styles)


def _pdf_answer(
    story: list[Any],
    row: Any,
    options: ExportOptions,
    styles: dict[str, ParagraphStyle],
) -> None:
    story.append(
        _pdf_paragraph(
            f"{_value(row, 'sequence', default=0)}. 正确答案：{_answer_text(row)}",
            styles["question"],
        )
    )
    explanation = _question_value(row, "explanation")
    if explanation:
        story.append(_pdf_paragraph(f"解析：{explanation}", styles["body"]))
    contents = _block_contents(row, "answer")
    if contents:
        for content in contents:
            units = (
                [unit for unit in content.units if unit.kind != "text"]
                if content.role.casefold() in ANSWER_ROLES
                else content.units
            )
            for unit in units:
                _pdf_unit(story, unit, options, styles)
    else:
        for unit in _fallback_units(row, "answer"):
            _pdf_unit(story, unit, options, styles)


def _render_pdf(
    paper: Any,
    variant: ExportVariant,
    options: ExportOptions,
    output: BinaryIO,
) -> None:
    styles = _pdf_styles()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=2.2 * cm,
        leftMargin=2.2 * cm,
        topMargin=(1.6 if variant == "answers" else 1.8) * cm,
        bottomMargin=(1.6 if variant == "answers" else 1.8) * cm,
    )
    story: list[Any] = [_pdf_paragraph(_paper_title(paper, variant), styles["title"])]
    if options.include_student_fields:
        info = Table(
            [["姓名：____________", "学号：____________", "班级：____________"]],
            colWidths=[5.53 * cm] * 3,
        )
        info.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#777777")),
                    ("FONTNAME", (0, 0), (-1, -1), PDF_FONT_NAME),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend([info, Spacer(1, 6)])
    story.append(_pdf_paragraph(_metadata_text(paper), styles["body"]))
    rows = _paper_questions(paper)
    if variant == "answer-sheet":
        for row in rows:
            story.append(
                _pdf_paragraph(
                    f"{_value(row, 'sequence', default=0)}. （{_format_number(_value(row, 'points', default=0))} 分）",
                    styles["question"],
                )
            )
            line_count = 3 if _question_value(row, "question_type", "questionType", default="") in {"简答题", "计算题"} else 1
            for _ in range(line_count):
                story.append(_pdf_paragraph("________________________________________________________________________", styles["body"]))
    else:
        previous_section = None
        for row in rows:
            section_title = str(_value(row, "section_title", "sectionTitle", default=""))
            if section_title and section_title != previous_section:
                story.append(_pdf_paragraph(section_title, styles["section"]))
                previous_section = section_title
            _pdf_question(story, row, options, styles)
    if variant == "answers":
        story.append(_pdf_paragraph("参考答案", styles["section"]))
        for row in rows:
            _pdf_answer(story, row, options, styles)

    def add_page_number(canvas, doc) -> None:
        canvas.saveState()
        canvas.setFont(PDF_FONT_NAME, 9)
        canvas.drawRightString(A4[0] - 2.2 * cm, 1.0 * cm, f"第 {doc.page} 页")
        canvas.restoreState()

    document.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
